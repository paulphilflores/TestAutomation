#This is a testing framework code for testing websites. 
#It will create a docx file with all the screenshots of the test steps of xlsx testscripts and generate xlsx test result summary.

#First step : Get test scripts and test data to be executed from testdata xlsx file.
##place all initial parameters here
import time
debug = True
err_message = "no error"
import openpyxl 
wb = openpyxl.load_workbook("testdata.xlsx")  #C:\Users\paulp\Downloads\py.projects\
ws = wb['td']

def startselenium():
    from selenium import webdriver
    web = webdriver.Edge(keep_alive=1)
    web.maximize_window()
    
    return web
web=startselenium()

def scriptdata(x_row): #get excel file and place in a phonebook scriptdata 
    
    if debug == True:
        x=1 #START FROM THE FIRST ROW OF WORKSHEET
        print('Total number of rows: '+str(ws.max_row)+'. And total number of columns: '+str(ws.max_column))
        while x < ws.max_row+1:
            values = [ws.cell(row=x,column=i).value for i in range(1,ws.max_column+1)]
            print(values)
            x+=1
        x=0

    #get excel file and place in a phonebook scriptdata 
    scriptdata={}
    y=1 #column for execute?[y/n]
    
    if ws.cell(row=x_row,column=y).value == "y":
        while y < ws.max_column+1:
            if ws.cell(row=x_row,column=y).value != "":
                scriptdata.update({ws.cell(row=1,column=y).value:ws.cell(row=x_row,column=y).value})
                y+=1
    else:
        scriptdata = "skip"    
                
    if debug == True:
        print(scriptdata)
        if scriptdata == "skip":
            pass
        else:
            print(scriptdata["testscripts"])

    return scriptdata

def getdata(param): #use this function to get specific test data from ws
    testdata = str(scriptdata[param]).replace("['","").replace("']","")
    return testdata

#Second step : Execute the step
class execute:
    strResult = ""
    err_message = ""
    action = ""
    object = ""
    objname = ""
    testdata = ""
   
    def executesteps(self):
        self.strResult="FAILED" 
        
        match self.action:
            case "open":
                if self.object == "browser":
                    web.get(self.testdata)
                    if web.current_url == self.testdata:
                        print(web.current_url)
                        self.strResult = "PASSED"
                        return self.strResult
                    else:
                        self.err_message = "web.current_url is not equal to "+testdata
                        return self.err_message
                    
            case "click":
                if self.object == "browser_button":
                    web.implicitly_wait(0.5)
                    browser_button = web.find_element(by='name', value=self.objname)
                    browser_button.click()
                    time.sleep(6)
                    self.strResult = "PASSED"
                    return self.strResult
                else:
                    self.err_message = "object is not defined in script, please add object"
                    return self.err_message
            
            case "sendkey":
                if self.object == "browser_textbox":
                    input = self.testdata
                    web.implicitly_wait(0.5)
                    textbox = web.find_element(by='name', value=self.objname)
                    textbox.send_keys(input)
                    if textbox.get_property("value")==input:
                        print(textbox.get_property("value"))                
                        self.strResult = "PASSED"
                        return self.strResult
                    else:
                        self.err_message = "textbox.get_property(""value"") is not equal to "+input
                        return self.err_message

            case "verify":
                match self.object:
                    case "text":
                        input = self.testdata
                        web.implicitly_wait(0.5)
                        match self.objname:
                            case "header_container":
                                obj = web.find_element(by='xpath',value='//*[@id="header_container"]/div[2]/span')
                            case "item_title_link": 
                                obj = web.find_element(by='xpath',value='//*[@id="item_*_title_link"]/div') 
                            case default:
                                self.err_message = "objname is not defined in code, please add objname"
                        
                        if obj.text==input: 
                            self.strResult = "PASSED"
                            return self.strResult
                        else: 
                            print(obj.text) 
                            self.err_message = "object text is not equal to "+input
                            return self.err_message
                    
                    case "buton":
                        input = self.testdata
                        web.implicitly_wait(0.5)
                        match self.objname:
                            case "header_container":
                                obj = web.find_element(by='xpath',value='//*[@id="header_container"]/div[2]/span')
                            case "item_title_link": 
                                obj = web.find_element(by='xpath',value='//*[@id="item_*_title_link"]/div') 
                            case default:
                                self.err_message = "objname is not defined in code, please add objname"
                        
                        if obj.text==input: 
                            self.strResult = "PASSED"
                            return self.strResult
                        else: 
                            print(obj.text) 
                            self.err_message = "object button is not equal to "+input
                            return self.err_message
                    
                    case default:
                       self.err_message = "object is not defined in code, please add object."
                       return self.err_message
            
            case default:
                self.err_message = "action is not defined in code, please add action."
                return self.err_message
   
#Third step : Take a screenshot of the executed step
def screenshot(ssname):
    import time
    timestamp = (time.strftime('%Y%M%H%M%S'))
    ss = ssname+"_"+timestamp+".png"
    import numpy as np 
    import cv2 
    import pyautogui 
    image = pyautogui.screenshot() 
    image = cv2.cvtColor(np.array(image),cv2.COLOR_RGB2BGR) 
    cv2.imwrite(ss, image)
    return ss
    #https://www.geeksforgeeks.org/how-to-take-screenshots-using-python/

#Fifth step : Place the screenshot to a document file, if the step fails, end the test and get a new test script.
def createscriptresultdocx(xrow,tsname,result,stepdesc,ssname):
    from docx import Document
    from docx.shared import Inches
    
    if xrow == 2: #for new script run instance
        document = Document()
        document.add_heading(tsname, 0)
    elif xrow > 2 and result == "FAILED":
        existingfile = tsname+"_"+"PASSED"+".docx"
        document = Document(existingfile)
    else:
        existingfile = tsname+"_"+result+".docx"
        document = Document(existingfile)
        
    document.add_heading(stepdesc, level=1)
    document.add_paragraph(result, style='Intense Quote')
    document.add_picture(ssname, width=Inches(5))
    document.add_page_break()
    document.save(tsname+"_"+result+'.docx')
    existingfile = tsname+"_"+result+".docx"
    print(existingfile)
    ##https://python-docx.readthedocs.io/en/latest/

#Six step: Create a summary of tests done


#This is the main script
starttime = time.time()
for x_row in range(2,ws.max_row+1):
    
    data = scriptdata(x_row)
    
    if data == "skip":
        pass
    else:
        filename = str(data["testscripts"]).replace("['","").replace("']","")
        print(filename)
        sheetname = "ts"
        import openpyxl 
        wb_ts = openpyxl.load_workbook(filename)  #C:\Users\paulp\Downloads\py.projects\
        ws_ts = wb_ts[sheetname]

        if debug == True:
            x=1 #initiate
            print('Total number of rows: '+str(ws_ts.max_row)+'. And total number of columns: '+str(ws_ts.max_column))
            while x < ws_ts.max_row+1:
                values = [ws_ts.cell(row=x,column=i).value for i in range(1,ws_ts.max_column+1)]
                print(values)
                x+=1
            x=0

        timestamp = (time.strftime('%Y%M%H%M%S'))
        tsname = filename.replace(".xlsx","")+"_"+timestamp
        xrow=2 #START IN ROW TWO OF TS WORKSHEET
        while xrow < ws_ts.max_row+1:  #loop for each script step
            print("xrow="+str(xrow))
            action = ws_ts.cell(row=xrow,column=2).value
            object = ws_ts.cell(row=xrow,column=3).value
            objname = ws_ts.cell(row=xrow,column=4).value
            if ws_ts.cell(row=xrow,column=5).value == None:
                pass
            else:
                param = ws_ts.cell(row=xrow,column=5).value
                testdata = str(data[param]).replace("['","").replace("']","")
            print(action +" "+ object +" "+  objname +" "+  testdata)
            
            stepexecute = execute()
            stepexecute.action = action
            stepexecute.object = object
            stepexecute.objname = objname
            stepexecute.testdata = testdata
            stepexecute.executesteps()
            result = stepexecute.strResult
            err_message = stepexecute.err_message
            #result = executesteps(action,object,objname,testdata)
            stepdesc = action +" "+ object +" "+  objname +"."
            ssname = screenshot(filename.replace(".xlsx",""))
            createscriptresultdocx(xrow,tsname,result,stepdesc,ssname)
            if result == "FAILED":
                print(err_message)
                
            xrow+=1
    #next script
    x_row+=1

print(x_row)
x_row=0

endtime=time.time()
print(endtime)
print(endtime-starttime)




##########################Helpful Codes#############################
#py -3 -m pip install openpyxl
#py -3 -m pip install selenium
#py -3 -m pip install numpy
#py -3 -m pip install pyautogui
#py -3 -m pip install pillow
#py -3 -m pip install opencv-python
#py -3 -m pip install python-docx

#https://ultimateqa.com/dummy-automation-websites/
#https://www.selenium.dev/documentation/webdriver/getting_started/first_script/
#https://python-docx.readthedocs.io/en/latest/
#https://www.askpython.com/python/examples/capture-screenshots
#https://docs.python.org/3/library/

##########################Function Library##########################

#openbrowser

#weight_kg = [81.65, 97.52, 95.25, 92.98, 86.18, 88.45]
#import numpy as np
## Create a numpy array np_weight_kg from weight_kg
#np_weight_kg = np.array(weight_kg)    
## Create np_weight_lbs from np_weight_kg
#np_weight_lb = np_weight_kg * 2.2
## Print out np_weight_lbs
#print(np_weight_lb)

#dict = {"country": ["Brazil", "Russia", "India", "China", "South Africa"],
#       "capital": ["Brasilia", "Moscow", "New Dehli", "Beijing", "Pretoria"],
#       "area": [8.516, 17.10, 3.286, 9.597, 1.221],
#       "population": [200.4, 143.5, 1252, 1357, 52.98] }
#import pandas as pd
#brics = pd.DataFrame(dict)
#print(brics)

#import re
#find_members = []
#for module in dir(re):
#    if module[:4]=="find":
#        find_members.append(module)
#print(find_members)    

## Create 2 new lists height and weight
#height = [1.87,  1.87, 1.82, 1.91, 1.90, 1.85]
#weight = [81.65, 97.52, 95.25, 92.98, 86.18, 88.45]
## Import the numpy package as np
#import numpy as np
##Create 2 numpy arrays from height and weight
#np_height = np.array(height)
#np_weight = np.array(weight)
#weight_kg = [81.65, 97.52, 95.25, 92.98, 86.18, 88.45]
#import numpy as np
##Create a numpy array np_weight_kg from weight_kg
#np_weight_kg = np.array(weight_kg)    
## Create np_weight_lbs from np_weight_kg
#np_weight_lbs = np_weight_kg*2.2
## Print out np_weight_lbs
#print(np_weight_lbs)

# edit the functions prototype and implementation
#def foo(a, b, c, *otherinput):
#    pass
#    if list(otherinput) == [4]:
#        return 1
#    if list(otherinput) == [4,5]:
#        return 2
#    
#def bar(a, b, c, **options):
###    pass
#    if options.get("magicnumber") == 6:
#        return False
#        
#    if options.get("magicnumber") == 7:
#        return True
# test code
#if foo(1, 2, 3, 4) == 1:
#    print("Good.")
#if foo(1, 2, 3, 4, 5) == 2:
#    print("Better.")
#if bar(1, 2, 3, magicnumber=6) == False:
#    print("Great.")
#if bar(1, 2, 3, magicnumber=7) == True:
#    print("Awesome!")

#a = set(["Jake", "John", "Eric"])
#b = set(["John", "Jill"])
#print(a.difference(b))
#print(a.union(b))
#print(b.symmetric_difference(a))
#print(b.intersection(a))

#import json
## fix this function, so it adds the given name
## and salary pair to salaries_json, and return it
#def add_employee(salaries_json, name, salary):
#    salaries = json.loads(salaries_json)
#    salaries[name] = salary
##    return json.dumps(salaries)
### test code
#salaries = '{"Alfred" : 300, "Jane" : 400 }'
#new_salaries = add_employee(salaries, "Me", 800)
#decoded_salaries = json.loads(new_salaries)
#print(decoded_salaries["Alfred"])
#print(decoded_salaries["Jane"])
#print(decoded_salaries["Me"])

#from functools import partial
#def func(u, v, w, x):
#    return u*4 + v*3 + w*2 + x
#
#p = partial(func,5,6,7)
#print(p(8))